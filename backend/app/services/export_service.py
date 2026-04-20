from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_pdf(
    filename: str,
    summary: str,
    highlights: list = None,
    chat_history: list = None
) -> BytesIO:
    """
    Generate a PDF file with summary, highlights, and chat history.
    
    Args:
        filename: Name of the original file
        summary: The summary text
        highlights: List of highlight dicts with 'selected_text' key
        chat_history: List of chat dicts with 'question' and 'answer' keys
        
    Returns:
        BytesIO object containing PDF data
    """
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a2e'),
        spaceAfter=12,
        alignment=1  # Center
    )
    elements.append(Paragraph(f"Study Notes: {filename}", title_style))
    elements.append(Spacer(1, 0.3*inch))
    
    # Metadata
    date_text = datetime.now().strftime("%B %d, %Y")
    elements.append(Paragraph(f"<b>Generated:</b> {date_text}", styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Summary Section
    elements.append(Paragraph("Summary", styles['Heading2']))
    elements.append(Paragraph(summary, styles['Normal']))
    elements.append(Spacer(1, 0.3*inch))
    
    # Highlights Section
    if highlights:
        elements.append(Paragraph("Highlights", styles['Heading2']))
        for idx, highlight in enumerate(highlights, 1):
            text = highlight.get('selected_text', '')
            elements.append(Paragraph(f"<b>{idx}.</b> {text}", styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
    
    # Chat History Section
    if chat_history:
        elements.append(Paragraph("Q&A History", styles['Heading2']))
        for idx, chat in enumerate(chat_history, 1):
            question = chat.get('question', '')
            answer = chat.get('answer', '')
            elements.append(Paragraph(f"<b>Q{idx}:</b> {question}", styles['Normal']))
            elements.append(Paragraph(f"<b>A{idx}:</b> {answer}", styles['Normal']))
            elements.append(Spacer(1, 0.1*inch))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer


def export_to_docx(
    filename: str,
    summary: str,
    highlights: list = None,
    chat_history: list = None
) -> BytesIO:
    """
    Generate a DOCX file with summary, highlights, and chat history.
    
    Args:
        filename: Name of the original file
        summary: The summary text
        highlights: List of highlight dicts with 'selected_text' key
        chat_history: List of chat dicts with 'question' and 'answer' keys
        
    Returns:
        BytesIO object containing DOCX data
    """
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Study Notes: {filename}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    date_text = datetime.now().strftime("%B %d, %Y")
    doc.add_paragraph(f"Generated: {date_text}").runs[0].font.size = Pt(10)
    doc.add_paragraph()  # Blank line
    
    # Summary Section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph()  # Blank line
    
    # Highlights Section
    if highlights:
        doc.add_heading("Highlights", level=1)
        for idx, highlight in enumerate(highlights, 1):
            text = highlight.get('selected_text', '')
            p = doc.add_paragraph(text, style='List Number')
        doc.add_paragraph()  # Blank line
    
    # Chat History Section
    if chat_history:
        doc.add_heading("Q&A History", level=1)
        for idx, chat in enumerate(chat_history, 1):
            question = chat.get('question', '')
            answer = chat.get('answer', '')
            
            q_para = doc.add_paragraph(f"Q{idx}: {question}")
            q_para.runs[0].bold = True
            
            a_para = doc.add_paragraph(f"A{idx}: {answer}")
            doc.add_paragraph()  # Blank line
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer